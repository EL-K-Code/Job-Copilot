from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.services import profile_onboarding
from app.services.profile_onboarding import (
    CVDocument,
    CVTextExtractionError,
    MAX_PROFILE_FACTS,
    ProfileExtraction,
    ProfileFactDraft,
    build_profile_memories,
    extract_cv_text,
    extract_profile_facts,
    manual_profile_to_review_rows,
    prepare_cv_documents,
    prepare_profile_extraction_input,
)


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_txt_cv_without_persisting_source_file():
    document = extract_cv_text(
        "candidate.txt",
        "Experience\nBuilt a Python data pipeline.\n".encode("utf-8"),
    )

    assert document.filename == "candidate.txt"
    assert "Built a Python data pipeline." in document.text


def test_extract_docx_includes_paragraphs_and_table_cells():
    document = Document()
    document.add_paragraph("Machine Learning Engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "SQL"
    buffer = BytesIO()
    document.save(buffer)

    extracted = extract_cv_text("candidate.docx", buffer.getvalue())

    assert "Machine Learning Engineer" in extracted.text
    assert "Python | SQL" in extracted.text


def test_rejects_unsupported_and_empty_cv_files():
    with pytest.raises(CVTextExtractionError, match="Unsupported file type"):
        extract_cv_text("candidate.rtf", b"profile")

    with pytest.raises(CVTextExtractionError, match="is empty"):
        extract_cv_text("candidate.txt", b"")


def test_prepare_cv_documents_enforces_file_count():
    uploads = [(f"cv-{index}.txt", b"Python") for index in range(6)]
    with pytest.raises(CVTextExtractionError, match="no more than 5"):
        prepare_cv_documents(uploads)


def test_profile_input_removes_only_exact_normalized_duplicate_lines():
    prepared = prepare_profile_extraction_input(
        [
            CVDocument(
                filename="cv-ai.txt",
                text="Python\nBuilt a data pipeline.\nShared achievement",
            ),
            CVDocument(
                filename="cv-data.txt",
                text=" python \nBuilt an analytics dashboard.\nShared achievement",
            ),
        ]
    )

    assert prepared.duplicate_lines_removed == 2
    assert prepared.unique_lines == 4
    assert prepared.text.casefold().count("python") == 1
    assert prepared.text.casefold().count("shared achievement") == 1
    assert "Built a data pipeline." in prepared.text
    assert "Built an analytics dashboard." in prepared.text
    assert "===== FILE: cv-ai.txt =====" in prepared.text
    assert "===== FILE: cv-data.txt =====" in prepared.text
    assert prepared.prompt_chars < prepared.original_chars + 100


def test_profile_input_does_not_fuzzy_delete_distinct_facts():
    prepared = prepare_profile_extraction_input(
        [
            CVDocument(filename="a.txt", text="Built a pipeline."),
            CVDocument(filename="b.txt", text="Built a data pipeline."),
        ]
    )

    assert prepared.duplicate_lines_removed == 0
    assert "Built a pipeline." in prepared.text
    assert "Built a data pipeline." in prepared.text


def test_profile_extraction_schema_caps_generated_fact_count():
    assert MAX_PROFILE_FACTS == 48
    facts = [
        ProfileFactDraft(
            type="skill",
            content=f"The candidate uses tool {index}.",
        )
        for index in range(MAX_PROFILE_FACTS + 1)
    ]
    with pytest.raises(ValueError):
        ProfileExtraction(facts=facts)


def test_extract_profile_facts_deduplicates_and_constrains_source(monkeypatch):
    captured_messages = []

    class FakeRunnable:
        def invoke(self, messages):
            captured_messages.extend(messages)
            return ProfileExtraction(
                facts=[
                    ProfileFactDraft(
                        type="project",
                        content="The candidate built a LangGraph workflow.",
                        source_file="cv.pdf",
                        topic="langgraph",
                        group_id="agentic_ai",
                    ),
                    ProfileFactDraft(
                        type="project",
                        content="The candidate built a LangGraph workflow.",
                        source_file="old-cv.pdf",
                        topic="langgraph",
                        group_id="agentic_ai",
                    ),
                    ProfileFactDraft(
                        type="skill",
                        content="The candidate uses Python.",
                        source_file="invented-name.pdf",
                        topic="python",
                        group_id="skills",
                    ),
                ]
            )

    monkeypatch.setattr(
        profile_onboarding,
        "get_structured_chat_model",
        lambda _schema: FakeRunnable(),
    )

    documents = [
        CVDocument(
            filename="cv.pdf",
            text="LangGraph and Python\nLangGraph and Python",
        )
    ]
    prepared = prepare_profile_extraction_input(documents)
    extraction = extract_profile_facts(documents, prepared_input=prepared)

    assert len(extraction.facts) == 2
    assert extraction.facts[0].source_file == "cv.pdf"
    assert extraction.facts[1].source_file == "CV import"
    assert captured_messages[-1].content.count("LangGraph and Python") == 1


def test_manual_profile_creates_one_review_row_per_non_empty_line():
    rows = manual_profile_to_review_rows(
        {
            "experience": "- Built a data pipeline.\n• Evaluated a model.\n",
            "skill": "Python\nSQL",
        }
    )

    assert [row["Category"] for row in rows] == [
        "experience",
        "experience",
        "skill",
        "skill",
    ]
    assert rows[0]["Fact"] == "Built a data pipeline."
    assert all(row["Source"] == "Manual profile" for row in rows)


def test_build_profile_memories_is_stable_deduplicated_and_user_verified():
    rows = [
        {
            "Use": True,
            "Category": "project",
            "Fact": "The candidate built a LangGraph workflow.",
            "Source": "cv.pdf",
            "Topic": "langgraph",
            "Group": "agentic_ai",
        },
        {
            "Use": True,
            "Category": "project",
            "Fact": "The candidate built a LangGraph workflow.",
            "Source": "another.pdf",
        },
        {
            "Use": False,
            "Category": "skill",
            "Fact": "The candidate uses Rust.",
            "Source": "cv.pdf",
        },
    ]

    first = build_profile_memories(rows)
    second = build_profile_memories(rows)

    assert first == second
    assert len(first) == 1
    assert first[0]["type"] == "project"
    assert first[0]["topic"] == "langgraph"
    assert first[0]["group_id"] == "agentic_ai"
    assert first[0]["source"] == "cv.pdf"
    assert first[0]["verified_by_user"] is True
    assert first[0]["id"].startswith("profile_project_")


def test_build_profile_memories_merges_existing_without_duplicate_content():
    existing = [
        {
            "id": "existing_python",
            "type": "skill",
            "topic": "python",
            "group_id": "skills",
            "content": "The candidate uses Python.",
        }
    ]
    rows = [
        {
            "Use": True,
            "Category": "skill",
            "Fact": "The candidate uses Python.",
            "Source": "new-cv.pdf",
        },
        {
            "Use": True,
            "Category": "skill",
            "Fact": "The candidate uses SQL.",
            "Source": "new-cv.pdf",
        },
    ]

    memories = build_profile_memories(rows, existing_memories=existing)

    assert len(memories) == 2
    assert memories[0]["id"] == "existing_python"
    assert memories[1]["content"] == "The candidate uses SQL."


def test_build_profile_memories_requires_one_kept_fact():
    with pytest.raises(ValueError, match="Keep at least one"):
        build_profile_memories(
            [
                {
                    "Use": False,
                    "Category": "skill",
                    "Fact": "The candidate uses Python.",
                }
            ]
        )
