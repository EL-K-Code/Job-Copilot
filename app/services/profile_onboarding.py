from __future__ import annotations

import io
import re
from collections.abc import Sequence
from hashlib import sha256
from pathlib import Path
from typing import Any, Literal

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel, Field, field_validator
from pypdf import PdfReader

from app.services.model_provider import get_structured_chat_model


MAX_CV_FILES = 5
MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_TOTAL_TEXT_CHARS = 120_000
PROFILE_TYPES = (
    "identity",
    "experience",
    "project",
    "education",
    "skill",
    "language",
    "certification",
    "achievement",
    "preference",
)
ProfileType = Literal[
    "identity",
    "experience",
    "project",
    "education",
    "skill",
    "language",
    "certification",
    "achievement",
    "preference",
]


PROFILE_EXTRACTION_SYSTEM_PROMPT = """
You extract a candidate profile from one or more CV documents for JobCopilot.

Return only facts explicitly supported by the supplied CV text. Every fact will be
shown to the user for approval before it becomes usable evidence.

Rules:
- Produce atomic facts: one independently verifiable statement per item.
- Never infer proficiency, ownership, scale, production use, leadership, dates,
  results, recency, or technologies that are not explicit in the CV.
- Preserve conservative wording. Prefer a narrower fact over a stronger paraphrase.
- Write each content field as a complete sentence beginning with "The candidate".
- Keep distinct technologies, responsibilities, outcomes, education records,
  languages, certifications, and projects as distinct facts when useful.
- Do not extract phone numbers, email addresses, postal addresses, nationality,
  marital status, age, photographs, or other unnecessary personal identifiers.
- source_file must be one of the filenames shown in the document headers. Use the
  filename containing the fact; when the same fact appears in several CVs, choose one.
- topic and group_id must be short lowercase snake_case audit labels.
- Remove duplicates across CV versions.

Return only structured data matching the requested schema.
""".strip()


class CVDocument(BaseModel):
    filename: str
    text: str


class ProfileFactDraft(BaseModel):
    type: ProfileType
    content: str = Field(min_length=1)
    source_file: str = ""
    topic: str = ""
    group_id: str = ""

    @field_validator("content", "source_file", "topic", "group_id")
    @classmethod
    def normalize_text(cls, value: str) -> str:
        return " ".join(str(value).strip().split())


class ProfileExtraction(BaseModel):
    facts: list[ProfileFactDraft] = Field(default_factory=list)


class CVTextExtractionError(ValueError):
    """Raised when an uploaded CV cannot provide usable text."""


def _clean_extracted_text(value: str) -> str:
    value = value.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    lines = [" ".join(line.split()) for line in value.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _extract_pdf_text(payload: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(payload))
        if reader.is_encrypted:
            raise CVTextExtractionError("Password-protected PDF files are not supported.")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except CVTextExtractionError:
        raise
    except Exception as exc:
        raise CVTextExtractionError("The PDF could not be read.") from exc


def _extract_docx_text(payload: bytes) -> str:
    try:
        document = Document(io.BytesIO(payload))
    except Exception as exc:
        raise CVTextExtractionError("The DOCX file could not be read.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_cv_text(filename: str, payload: bytes) -> CVDocument:
    """Extract text locally from one supported CV without persisting the source file."""
    normalized_name = Path(filename).name.strip()
    if not normalized_name:
        raise CVTextExtractionError("Every uploaded CV needs a filename.")
    if not payload:
        raise CVTextExtractionError(f"{normalized_name} is empty.")
    if len(payload) > MAX_FILE_BYTES:
        raise CVTextExtractionError(
            f"{normalized_name} is larger than the 10 MB per-file limit."
        )

    suffix = Path(normalized_name).suffix.casefold()
    if suffix == ".pdf":
        text = _extract_pdf_text(payload)
    elif suffix == ".docx":
        text = _extract_docx_text(payload)
    elif suffix == ".txt":
        text = payload.decode("utf-8-sig", errors="replace")
    else:
        raise CVTextExtractionError(
            f"Unsupported file type for {normalized_name}. Use PDF, DOCX or TXT."
        )

    cleaned = _clean_extracted_text(text)
    if not cleaned:
        if suffix == ".pdf":
            raise CVTextExtractionError(
                f"No selectable text was found in {normalized_name}. "
                "Scanned image-only PDFs are not supported yet; use DOCX, TXT, or a text PDF."
            )
        raise CVTextExtractionError(f"No usable text was found in {normalized_name}.")
    return CVDocument(filename=normalized_name, text=cleaned)


def prepare_cv_documents(
    uploads: Sequence[tuple[str, bytes]],
) -> list[CVDocument]:
    if not uploads:
        raise CVTextExtractionError("Upload at least one CV.")
    if len(uploads) > MAX_CV_FILES:
        raise CVTextExtractionError(f"Upload no more than {MAX_CV_FILES} CV files at once.")

    documents = [extract_cv_text(filename, payload) for filename, payload in uploads]
    total_chars = sum(len(document.text) for document in documents)
    if total_chars > MAX_TOTAL_TEXT_CHARS:
        raise CVTextExtractionError(
            "The combined CV text is too long. Upload fewer or shorter CV versions."
        )
    return documents


def extract_profile_facts(documents: Sequence[CVDocument]) -> ProfileExtraction:
    if not documents:
        raise ValueError("At least one extracted CV document is required.")

    source_names = {document.filename for document in documents}
    document_text = "\n\n".join(
        f"===== FILE: {document.filename} =====\n{document.text}"
        for document in documents
    )
    structured_llm = get_structured_chat_model(ProfileExtraction)
    result = structured_llm.invoke(
        [
            SystemMessage(content=PROFILE_EXTRACTION_SYSTEM_PROMPT),
            HumanMessage(content=f"CV documents:\n\n{document_text}"),
        ]
    )

    unique_facts: list[ProfileFactDraft] = []
    seen: set[tuple[str, str]] = set()
    for fact in result.facts:
        normalized_content = " ".join(fact.content.split())
        key = (fact.type, normalized_content.casefold())
        if not normalized_content or key in seen:
            continue
        source_file = fact.source_file if fact.source_file in source_names else "CV import"
        unique_facts.append(
            fact.model_copy(
                update={
                    "content": normalized_content,
                    "source_file": source_file,
                }
            )
        )
        seen.add(key)
    return ProfileExtraction(facts=unique_facts)


def extraction_to_review_rows(extraction: ProfileExtraction) -> list[dict[str, Any]]:
    return [
        {
            "Use": True,
            "Category": fact.type,
            "Fact": fact.content,
            "Source": fact.source_file or "CV import",
            "Topic": fact.topic,
            "Group": fact.group_id,
        }
        for fact in extraction.facts
    ]


def _split_manual_facts(value: str) -> list[str]:
    facts: list[str] = []
    for line in str(value).splitlines():
        normalized = re.sub(r"^[\s\-•*\d.)]+", "", line).strip()
        if normalized:
            facts.append(" ".join(normalized.split()))
    return facts


def manual_profile_to_review_rows(sections: dict[str, str]) -> list[dict[str, Any]]:
    """Convert guided one-fact-per-line fields into the common human review table."""
    rows: list[dict[str, Any]] = []
    for category, value in sections.items():
        if category not in PROFILE_TYPES:
            raise ValueError(f"Unsupported manual profile category: {category}")
        for fact in _split_manual_facts(value):
            rows.append(
                {
                    "Use": True,
                    "Category": category,
                    "Fact": fact,
                    "Source": "Manual profile",
                    "Topic": "",
                    "Group": "",
                }
            )
    return rows


def _slug(value: str, *, fallback: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "_", value.casefold()).strip("_")
    return (normalized[:64].rstrip("_") or fallback)


def _stable_memory_id(category: str, content: str) -> str:
    digest = sha256(f"{category}\n{content.casefold()}".encode("utf-8")).hexdigest()[:12]
    return f"profile_{category}_{digest}"


def build_profile_memories(
    review_rows: Sequence[dict[str, Any]],
    *,
    existing_memories: Sequence[dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    """Build validated atomic memories only from rows explicitly kept by the user."""
    memories = [dict(memory) for memory in (existing_memories or [])]
    seen_content = {
        str(memory.get("content", "")).strip().casefold()
        for memory in memories
        if str(memory.get("content", "")).strip()
    }
    seen_ids = {
        str(memory.get("id", "")).strip()
        for memory in memories
        if str(memory.get("id", "")).strip()
    }

    for row in review_rows:
        if not bool(row.get("Use", True)):
            continue
        category = str(row.get("Category", "")).strip().casefold()
        content = " ".join(str(row.get("Fact", "")).strip().split())
        if category not in PROFILE_TYPES:
            raise ValueError(f"Unsupported profile category: {category or 'empty'}")
        if not content:
            raise ValueError("Every selected profile fact must contain text.")
        if content.casefold() in seen_content:
            continue

        topic = _slug(str(row.get("Topic", "")) or content, fallback=category)
        group_id = _slug(
            str(row.get("Group", "")) or f"{category}_profile",
            fallback=f"{category}_profile",
        )
        memory_id = _stable_memory_id(category, content)
        suffix = 2
        while memory_id in seen_ids:
            memory_id = f"{_stable_memory_id(category, content)}_{suffix}"
            suffix += 1

        memory: dict[str, Any] = {
            "id": memory_id,
            "type": category,
            "topic": topic,
            "group_id": group_id,
            "content": content,
            "verified_by_user": True,
        }
        source = " ".join(str(row.get("Source", "")).strip().split())
        if source:
            memory["source"] = source
        memories.append(memory)
        seen_content.add(content.casefold())
        seen_ids.add(memory_id)

    if not memories:
        raise ValueError("Keep at least one verified profile fact before activation.")
    return memories


def memories_to_review_rows(memories: Sequence[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "Use": True,
            "Category": str(memory.get("type", "identity")),
            "Fact": str(memory.get("content", "")),
            "Source": str(memory.get("source", "JobCopilot backup")),
            "Topic": str(memory.get("topic", "")),
            "Group": str(memory.get("group_id", "")),
        }
        for memory in memories
    ]
